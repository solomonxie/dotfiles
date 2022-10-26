import pandas as pd
from docx import Document
from docx.shared import Inches


def main():
    df = pd.read_csv('~/Downloads/bible/kjv.tsv', delimiter='\t', header=None, names=[
        'book', 'alias', 'bn', 'cn', 'vn', 'verse',
    ])
    __import__('pudb').set_trace()
    for (bname, bn, alias), book in df.groupby(['book', 'bn', 'alias'], sort=False):
        doc = Document()
        doc.add_heading(bname, 0)
        for cn, chapter in book.groupby(['cn'], sort=False):
            # table = doc.add_table(rows=1, cols=3)
            doc.add_heading(f'{bname} {cn}', level=1)
            table = doc.add_table(rows=1, cols=1)
            c = table.rows[0].cells[0]
            c.width = Inches(0.5)
            # p = c.add_paragraph('')
            for _, v in chapter.iterrows():
                p = c.add_paragraph('')
                t_vn = p.add_run(f'{v.vn} ')
                t_vn.italic = True
                t_verse = p.add_run(f'{v.verse} ')
                t_verse.bold = False
            doc.add_page_break()
            print(f'Added Book {bname} - {cn}')
            # break
        doc.save(f'/tmp/kjv_from_tsv_{bn}_{bname}.docx')
        break
    doc.save('/tmp/kjv_from_tsv.docx')
    return


if __name__ == '__main__':
    main()
